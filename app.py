import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import datetime
from supabase import create_client, Client

# --- 1. CONEXIÓN AL MOTOR SUPABASE ---
# Streamlit buscará estas llaves en la sección de "Secrets"
try:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    supabase: Client = create_client(url, key)
except:
    st.error("⚠️ Falta configurar las llaves de Supabase en los secretos de la página.")

# --- 2. CONFIGURACIÓN VISUAL ---
st.set_page_config(page_title="Profe.Educa", layout="wide", page_icon="🍎")
st.sidebar.title("🍎 Profe.Educa")
menu = st.sidebar.radio("Menú Principal", ["Inicio", "Planeación Semanal", "Texto Reflexivo Diario", "Evaluación Trimestral", "Admin"])

# --- 3. MÓDULO: PLANEACIÓN ---
if menu == "Planeación Semanal":
    st.title("📋 Planeación de Trayectos")
    with st.form("form_p"):
        ec = st.text_input("Educador Comunitario")
        eca = st.text_input("E.C. de Acompañamiento")
        meta = st.text_area("Meta de la semana")
        if st.form_submit_button("Generar Word"):
            doc = Document()
            doc.add_heading('PLANEACIÓN CONAFE', 0)
            doc.add_paragraph(f"EC: {ec}\nMeta: {meta}")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Descargar", buffer, "Planeacion.docx")

# --- 4. MÓDULO: REFLEXIÓN DIARIA (GUARDA EN SUPABASE) ---
elif menu == "Texto Reflexivo Diario":
    st.title("✍️ Registro de Relación Tutora")
    with st.form("registro"):
        alumno = st.text_input("Nombre del Alumno")
        notas = st.text_area("Anotaciones del día")
        if st.form_submit_button("Guardar en la Nube"):
            data = {"alumno_nombre": alumno, "contenido_reflexivo": notas}
            supabase.table("reflexiones").insert(data).execute()
            st.success("✅ Guardado permanentemente en Supabase")

# --- 5. MÓDULO: EVALUACIÓN TRIMESTRAL ---
elif menu == "Evaluación Trimestral":
    st.title("📊 Resumen del Periodo")
    busqueda = st.text_input("Nombre del alumno a evaluar")
    if st.button("Jalar datos de la base"):
        res = supabase.table("reflexiones").select("*").eq("alumno_nombre", busqueda).execute()
        if res.data:
            st.write(f"Encontrados {len(res.data)} registros diarios.")
            st.dataframe(pd.DataFrame(res.data))
